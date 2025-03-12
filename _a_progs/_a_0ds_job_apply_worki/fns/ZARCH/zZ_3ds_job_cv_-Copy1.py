#%%writefile cv.py

#
def ds_cv():
    direc_exp = '/Users/yerik/_apple_source/d_OUT/_0ds/_0_ds_Employment/'
    ####################################################
    def get_multi_line_input():
        print("\n\n*** START : PASTE your job description.\nTERMINATE : type  * 'jd' * on a new line to finish: \n\n*------------>")
        lines = []
        while True:
            line = input()
            if line == 'jd':
                break
            lines.append(line)
        return '\n'.join(lines)

    def save_text_to_file(text, filename):
        with open(filename, 'w') as file:
            file.write(text)

    job_description = get_multi_line_input()
    if job_description:
        filename = f"{direc_exp}job_description.txt"
        save_text_to_file(job_description, filename)
        print(f"Job description saved to {filename}")
    else:
        print("No job description provided.")

    ####################################################
    def read_file_to_variable(filename):
        try:
            with open(filename, 'r') as file:
                content = file.read()
                return content
        except FileNotFoundError:
            return "File not found. Please check the file name and try again."
        except Exception as e:
            return f"An error occurred: {e}"

        # Replace 'job_description.txt' with your actual file name if different
    filename = f"{direc_exp}job_description.txt"
    job_desc = read_file_to_variable(filename)
    
    if job_desc:
        print("Job description loaded successfully.")
        # You can now use the job_desc variable as needed
    else:
        print("No job description found in the file.")

    ####################################################







    
    import calendar
    from docx2pdf import convert
    from datetime import date
    import docx
    today = date.today();date = today.strftime("%d_%m_%Y")
    year = f"{date[6:10]}"; month_number=f"{date[3:5]}"; month=calendar.month_name[int(month_number)]; day=f"{date[0:2]}"
    written_date = f"{month} {day}, {year} ";day_name=calendar.day_name[today.weekday()] 
    import subprocess
    from docx import Document
    import pyperclip
    document = Document()

    ##--------------------------------------------------VARIABLES-------------------------------------->
    pos = input('position name -- ds/ml/dl::: <E> for CUSTOM position ... ')
    #matc = input('matched position -- ds/ml/dl::: ')

    if pos == "ds":
        position = "Data Scientist"
        matched_role = position
    elif pos == "ml":
        position = "Machine Learning Specialist"
        matched_role = position
    elif pos == "dl":
        position = "Deep Learning Specialist"
        matched_role = position
    else:
        position = input('What is the position name? ...')
        matched_role = position
    print(position)


    #If statement for matched trigger words ...

    if pos == "ds":#--------------------------------------------------------------DS

        trigger_word = "Statistics and data manipulation"
        #trigger expertises
        trigger_expertise_1="optimization"
        trigger_expertise_2="forecast"
        trigger_expertise_3="analysis"
        #trigger sentence
        trigger_sentence = 'The science of Statistics is one thing I am amazed by.'
    elif pos == "ml":#--------------------------------------------------------------ML

        trigger_word = "Machine Learning methods"
        #trigger expertises
        trigger_expertise_1="optimization"
        trigger_expertise_2="forecast"
        trigger_expertise_3="analysis"
        #trigger sentence
        trigger_sentence = 'The science of Statistics is one thing I am amazed by.'

    elif pos == "dl":#-------------------------------------------------------------DL

        trigger_word = 'Artificial Intelligence and Supervised Learning'
        #trigger expertises
        trigger_expertise_1="optimization"
        trigger_expertise_2="forecast"
        trigger_expertise_3="analysis"
        #trigger sentence
        trigger_sentence = 'The science of Statistics is one thing I am amazed by.'

    else:#--------------------------------------------------------------------CUSTOM   
        #trigger word
        #---> {trigger_word} has been on the surface of my work as a {matched_role} more than ever before.
        #---> since then I have been involved in numerous projects involving {trigger_word}::: ''')
        trigger_word = 'Analizing Data, Buidling python programs, EDA'
        #trigger expertises
        #---> opportunities where I can use my {trigger_expertise_1}, {trigger_expertise_2}, and {trigger_expertise_3} expertises
        trigger_expertise_1="optimization"
        trigger_expertise_2="forecast"
        trigger_expertise_3="analysis"
        #trigger sentence
        #---> responsability of this job: ... I've noticed you are looking for a {position} to {trigger_sentence}::: ''')
        trigger_sentence = 'The science of Statistics is one thing I am amazed by.'


    #li_res = input('LINKED IN COVER')


    hirm_name = 'Hiring Manager'
    position_learn_from ="Linked-In"
    company_name= input(" Company Name : ")



    #EDIT SCRIPT SAY ONE OF MY MOST similiar projects then this and its description

    skillset = '''Machine Learning Methods
    Artificial Neural Network
    Computer Vision
    Feature Extraction
    Natural Language Processing
    Supervised Learning
    Deep Learning
    Machine Learning Tools
    Keras
    Amazon SageMaker
    SAS
    MATLAB
    TensorFlow
    PyTorch
    pandas
    Machine Learning Deliverables
    Model Optimization
    Machine Learning Model
    Machine Learning
    Data Science
    Data Science Consultation
    Machine Learning Techniques
    Anomaly Detection
    Classification
    Logistic Regression
    Linear Regression
    Principal Component Analysis
    Decision Tree
    Machine Learning Languages
    R
    Python
    Other skills
    Statistical ProgrammingData ProcessingData Analytics & Visualization Software '''

    #-----------------------------------------------------------------------------------------
    # #Create Folder
    # path_company_arch = f'{path_source_arch}/Jobs/Job_Applications/{company_name}/{position}' #Create folder in Archive for Job
    # if not os.path.exists(path_company_arch):
    #     os.makedirs(path_company_arch)

    #-----------------------------------------------------------------------------------------------------------------
    #-------------------------------------------------Cover-Letter----------------------------------------------------------------
    #-----------------------------------------------------------------------------------------------------------------
    #-----------------------------------------------------------------------------------------------------#### 1
    #### 1 Give a killer headline: When in doubt, use: Trgger word + Adjective(size, texture, shape , feel) +keyword +Promise
    killer_Headline = f"""I was excited to find this job opportunity is still open! {trigger_word} has been on the surface 
    of my work as a {matched_role} more than ever before."""
    #-----\
    #-----/
    #f"I’ve been passionate about 
    #f"I was excited to learn of this job opportunity "
    #INSTRUCTOR# f“As a teacher, I believe every student deserves the opportunity to learn"  #INSTRUCTOR
    #-----\
    #-----/
    #-----------------------------------------------------------------------------------------------------#### 2
    #### 2 Personalized_Greetings: Research the name of HR manager or use Dear Sir/Madam
    Personalized_Greetings= f'''Dear Sir, I thank you for your time in reading this letter. I would love to introduce myself, 
    and related work experience that show my dedication and my hard work as {matched_role}. 

    '''
    #-----\
    #-----/
    #UPWORK#
    #JOB# f'''Dear Sir, I thank you for your time in reading this letter and I would love to let you know how I can help {company_name} {main_task_verb}'''
    #JOB# f'''Dear Sir {hirm_name} It is with great enthusiasm that I write to you in response to the job posting Data Scientist. I learned of this position through Linked In '''
    #-----\
    #-----/
    #-----------------------------------------------------------------------------------------------------#### 3
    #### 3 Introduce Yourself:Open Strong and compliment the company
    P1_Introduce_yourself = f'''My name is Yeriko Vargas, I became a Statistician in May 2019 and since then 
    I have been involved in numerous projects involving {trigger_word}. I have decided to keep finding opportunities 
    where I can use my {trigger_expertise_1}, {trigger_expertise_2}, and {trigger_expertise_3} expertise. As I am 
    aware of your needs, I had to reach out and offer my services. I have 8+ years of involvement with statistics work and 5+ artificial neural network methodologies. '''
    #-----\
    #-----/

    #-----\
    #-----/
    #-----------------------------------------------------------------------------------------------------#### 4
    #### 4 Why you are a great fit for the company:Write a short Summaryof your career and skills, and tailor it to fit to company
    P2_Why_you_great_fit= f'''The opportunity of {position} at {company_name}, would be best if taken into my hands. 
    As {matched_role} I have had the opportunity to work with big data 
    by handling millions of rows, starting with EDA analysis till the deployment stage of Machine Learning and Deep 
    Learning models. 
    Aside from my 5+ years' experience in python and R programming skills, I have a powerful base of ML techniques and its 
    algorithms. These strong bases of math and statistics were transferred in my B.S. and M.S. in Applied Statistics from 
    teachers and advisors at Oakland University; in particular Dr. Ravindra Khattree. After much work in math and statistics, 
    I gave the most important switch in my career by deciding to do Deep Learning research. This decision took me to understand 
    the deep layers in neural networks, and how the logic of the perceptron theorem work. The main area of research was 
    {trigger_expertise_1}, and specifically what activation functions are best fitting the neural network models when 
    {trigger_expertise_1}, and {trigger_expertise_2} are needed. After creating so much momentum with this opportunity 
    I was taken into a path for similar projects working full time as {matched_role} for Ford, FCA, and 
    as a AI/ML Researcher for an AI consulting company; Soothsayer Analytics. 
    '''
    #-----\
    #-----/
    #f'''similarities within the duties and responsabilities for the {position} job description at {company_name} {position} needsat 
    #-----\
    #-----/
    #### 5 Why the COMPANY is  agreat fit for you:What excites you about working there? what do you want to learn there?
    P3_Why_company_great_fit = f'''I was so lucky to find the {position} position, I would really be a great {matched_role} 
    at {company_name}. I know that {trigger_word} projects are a great fit for me. Please consider taking the next step. 
    I want to make sure to tell you that I am huge in organization, very self-driven and really creative; when it comes 
    to data, I believe every data set can be considered as a story to tell. I have the ability to tackle enormous 
    {trigger_word} problems. I have experienced {trigger_word} work in detailed and I have gained so much understanding 
    around this topic, and I can assure you I have what it takes to help as {matched_role}. '''

    #### 6 Finish strong and stay in touch:Repeat why you are a great fit, explain how and when you are going to contact them
    Closing_stay_in_touch = f'''Thank you I believe that there is so much I could passionately talk 
    about other related experiences to the work needed here. I think that the 
    purpose of the {position} at {company_name} closely matches my career goals as well. I really would 
    be so happy and excited to have the chance to speak with someone and learn more about this position, as 
    I want to find a deal to work here at {company_name} as for the {position} position'''

    #### 7 All Best...
    All_the_best_name = f'''All best,
    Yeriko Vargas'''

    killer_Headline=killer_Headline.replace('\n', '')
    Personalized_Greetings = Personalized_Greetings.replace('\n', '')
    P1_Introduce_yourself = P1_Introduce_yourself.replace('\n', '')
    P2_Why_you_great_fit = P2_Why_you_great_fit.replace('\n', '')
    P3_Why_company_great_fit =P3_Why_company_great_fit.replace('\n', '')
    Closing_stay_in_touch = Closing_stay_in_touch.replace('\n', '')

    #-----------------------------------------------------------------------------------------------------------------

    #Write txt document
    cv_text =f'''Yeriko Vargas 
    {matched_role} 
    ygvargas93@gmail.com | in/yeriko-vargas/ | twitter.com/vargasyeriko | (646)-771-6111

    {written_date}

    {company_name}
    re: {position}

    {killer_Headline}

    {Personalized_Greetings} {P1_Introduce_yourself}

    {P2_Why_you_great_fit}

    {P3_Why_company_great_fit}

    {Closing_stay_in_touch}

    {All_the_best_name}'''

    # Write file
    document.add_paragraph(f'{cv_text}')




            #Write txt document
    cv_text_chat_gpt =f'''
    HEY start here reading my CV until the end :
    
    
    Yeriko Vargas 
    {matched_role} 
    ygvargas93@gmail.com | in/yeriko-vargas/ | twitter.com/vargasyeriko | (646)-771-6111

    {written_date}

    {company_name}
    re: {position}

    {killer_Headline}

    {Personalized_Greetings} {P1_Introduce_yourself}

    {P2_Why_you_great_fit}

    {P3_Why_company_great_fit}

    {Closing_stay_in_touch}

    {All_the_best_name}
    
    
    _________________________ END OF COVER LETTER... 
    
    * NEXT STATEMENT ARE MY INSTRUCTIONS FOR YOU :::
    
    (1) : NOW you CHAT GPT change a little bit in my CV very nicely (flowy structure making sure to have by itself the CV in a 
    first markdown )matching the job description. Also mathc most important skills in experiences also found below.
    . make sure
    you use a very neutral tone. do not adjust so much the way I say things, but greatly improve my chances
    to get a match from the company I am applying to ability to match job descriptions with applicants. I do not 
    want my application to be denied so please do your best to check in with the most important words we need to 
    highlight you know. ALSO AT THE VERY END in a separate markdown provide me with those keywords you found,
    that will greatly improve the odds of me getting the position. 
    :::::
    ------- END OF INSTRUCTIONS 
    
    * NEXT STATEMENT ARE MY EXPERIENCES ::: (a) trhough 
    

    (a) : at CAI, my role as an ML/AI Data Scientist involved developing sophisticated financial forecasting algorithms 
    using Python and Deep Learning. I leveraged Large Language Models for sentiment analysis of tweets and news, 
    enhancing price predictions. I built efficient Python pipelines incorporating TensorFlow and scikit-learn for 
    audio data analytics and NLP. A key contribution was the creation of a Voice Assistant for podcast episode 
    summaries and recommendations, highlighting my skills in audio categorization and premium revenue analytics.

    
    (b) : As Lead Data Scientist at FORD, I managed large datasets from the CVDOS database, utilizing Python 
    and AWS SageMaker. My responsibilities included data transformation, visualization, and daily statistical 
    reporting. I collaborated on machine learning projects for identifying car defects using Cluster Analysis 
    and PCA methodology, and automated data transfer processes, reducing human error and enhancing efficiency

    
    (c) : At Soothsayer Analytics, I focused on applying Machine Learning to fraud detection and customer segmentation.
    I used Python for data mining and visualization, improving model accuracy with SHAP and Random Forest models. 
    My role involved anomaly detection in service shops across 500 locations, using tools like Tableau, matplotlib, 
    and Power BI for effective visualization. 
    
    (d) : At Official Driving School, I leveraged data visualization tools like pivot tables and ggplot2
    in Rstudio to create dynamic business dashboards. My role included developing R-Keras machine learning
    models to forecast sales and optimize marketing strategies, blending business intelligence with AI-driven
    approaches for enhanced outcomes

    
    (e) : My role as a Deep Learning Analyst at Fiat Chrysler Automotive involved developing models using 
    R-Keras and Python to forecast profitable vehicle configurations. I analyzed data from half a million
    vehicles, applying Neural Networks to optimize automotive storage, significantly saving costs by 
    identifying optimal car configurations
    
    ]
    ------- end of my experiences
    
    
    * NEXT STATEMENT is JOB DESCRIPTION :::

    {job_desc}

    REMEMBER ! hey do not exaggerate things, make it relatable but make sure you 
    understand I am not lying right here , just relate 
    '''

        # Write file
    





        # Save File
        #document.save(f'test.docx')
    pyperclip.copy(cv_text_chat_gpt)



    print(" YOUR CV is BEING WRITTEN and NOW GO to !!!! CHAT GPT and paste !!! return to FINISH CV... \n")

# 4 --<<EXPORTS>>
    direc_exp = '/Users/yerik/_apple_source/d_OUT/_0ds/_0_ds_Employment/'
    direc_apply_now = '/Users/yerik/_apple_source/d_OUT/_0ds/apply_now/'

    
    document.save(f'{direc_exp}Yeriko Vargas - Cover Letter.docx')#WORD
    # OPEN CV
    subprocess.Popen(['open', f'{direc_exp}Yeriko Vargas - Cover Letter.docx'])
    # EDIT then come back and 
    input('BEFORE ::: <ENTER> ::: (1) convert to pdf & SENT to apply now !!! ... OPENING CV (i) modify & (ii) save (iii) come back' )
    
    # (1)2PDF
    convert(f"{direc_exp}Yeriko Vargas - Cover Letter.docx",
            f"{direc_apply_now}Yeriko Vargas-Cover_Letter - {company_name}_{position}-{date}.pdf" ) #pdf
    input("\nENTER to PASTE CV to DESKTOP \n")
    convert(f"{direc_exp}Yeriko Vargas - Cover Letter.docx",
            f"/Users/yerik/Desktop/Yeriko Vargas - Cover Letter.pdf") 
    # (2)zARCHIVE
    # direc_arch = '/Users/yerik/_apple_source/d_OUT/zARCHIVE/_0ds/apply_now/cvs/'
    # my_cv = docx.Document(f'{direc_exp}Yeriko Vargas-Cover_Letter - {company_name}_{position}-{date}.docx') # read modified
    # my_cv.save(f'{direc_arch}Yeriko Vargas-Cover_Letter - {company_name}_{position}-{date}.docx')#WORD
    #print("SAVED TO ARCHIVE")

    # ARCHIVE
    # OPEN CV again 
    pyperclip.copy(job_desc)
    input("::: ATTN -> JD COPIED ::: ENTER TO OPEN DOC , paste JD to ARCHIVE and ENTER ....")

    subprocess.Popen(['open', f'{direc_exp}Yeriko Vargas - Cover Letter.docx'])
    input('')
    convert(f"{direc_exp}Yeriko Vargas - Cover Letter.docx",
            f"{direc_apply_now}applied_{company_name}_{date}_Yeriko Vargas-Cover_Letter - {company_name}_{position}-.pdf" )





    
    return document


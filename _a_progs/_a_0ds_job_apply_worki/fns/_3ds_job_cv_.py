#%%writefile cv.py
import sys
import shutil
#
def ds_cv():

    direc_exp = '/Users/yerik/_apple_source/all_data/temp/_0ds/_0_ds_Employment/'
    ####################################################
    def get_multi_line_input():
        print("\n\n*** START : PASTE your job description.\nTERMINATE : type  * 'jd' * on a new line to finish: \n\n*------------>")
        lines = []
        while True:
            line = input()
            if line == 'jd':
                break
            lines.append(line)
        return '\n'.join(lines)

    def save_text_to_file(text, filename):
        with open(filename, 'w') as file:
            file.write(text)

    job_description = get_multi_line_input()
    if job_description:
        filename = f"{direc_exp}job_description.txt"
        save_text_to_file(job_description, filename)
        print(f"Job description saved to {filename}")
    else:
        print("No job description provided.")

    ####################################################
    def read_file_to_variable(filename):
        try:
            with open(filename, 'r') as file:
                content = file.read()
                return content
        except FileNotFoundError:
            return "File not found. Please check the file name and try again."
        except Exception as e:
            return f"An error occurred: {e}"

        # Replace 'job_description.txt' with your actual file name if different
    filename = f"{direc_exp}job_description.txt"
    job_desc = read_file_to_variable(filename)
    
    if job_desc:
        print("Job description loaded successfully.")
        # You can now use the job_desc variable as needed
    else:
        print("No job description found in the file.")

    ####################################################
    ##############################################RESUME
    ####################################################

    user_decision = 'y'#input("\n to OPEN : MATCHED - RESUME - descriptions : - do CLIPBOARD next ::: <y/E> :::\n").lower()
    if user_decision == 'y':

        
        res_path = '/Users/yerik/_apple_source/PY/functions/_pyENVS/_Source/src_env0/worki/fns/sub_fns'
        sys.path.append(res_path)  # Temporarily add res_path to the Python import path
        from _0ds_resume_builder import resume  # Import the function
        schema ,cai_e1 ,ford_e2,added_e3= resume(job_desc)  # Call the function and capture its return value
        
        #print(f'{schema}','\n\n\n\n\n') ######################## match_desc *******************
        #print(f'{cai_e1}')
        #print(f'{ford_e2}')
        
        sys.path.remove(res_path) 


        # res_path = '/Users/yerik/_apple_source/PY/functions/_pyENVS/_Source/src_env0/worki/fns/sub_fns'
        # exec(open(f"{res_path}/_0ds_resume_builder.py", encoding="utf-8").read()) 
        # print(f'{schema_1}')
        
    elif user_decision == '':
        print("Not opening another function.")
    else:
        print("Invalid input. Exiting.")


    ####################################################
    ####################################################
    ####################################################
    


    
    import calendar
    from docx2pdf import convert
    from datetime import date
    import docx
    today = date.today();date = today.strftime("%d_%m_%Y")
    year = f"{date[6:10]}"; month_number=f"{date[3:5]}"; month=calendar.month_name[int(month_number)]; day=f"{date[0:2]}"
    written_date = f"{month} {day}, {year} ";day_name=calendar.day_name[today.weekday()] 
    import subprocess
    from docx import Document
    import pyperclip
    document = Document()

    ##--------------------------------------------------VARIABLES-------------------------------------->
    pos = input('\n\n POSITION name - or - ds/ml/dl::: <E> for CUSTOM position ... ****************\n')
    #matc = input('matched position -- ds/ml/dl::: ')

    if pos == "ds":
        position = "Data Scientist"
        matched_role = position
    elif pos == "ml":
        position = "Machine Learning Specialist"
        matched_role = position
    elif pos == "dl":
        position = "Deep Learning Specialist"
        matched_role = position
    else:
        
        position = input('\nWhat is the position name? ...')
        matched_role = position
    print(position)


    #If statement for matched trigger words ...

    if pos == "ds":#--------------------------------------------------------------DS

        trigger_word = "Statistics and data manipulation"
        #trigger expertises
        trigger_expertise_1="optimization"
        trigger_expertise_2="forecast"
        trigger_expertise_3="analysis"
        #trigger sentence
        trigger_sentence = 'The science of Statistics is one thing I am amazed by.'
    elif pos == "ml":#--------------------------------------------------------------ML

        trigger_word = "Machine Learning methods"
        #trigger expertises
        trigger_expertise_1="optimization"
        trigger_expertise_2="forecast"
        trigger_expertise_3="analysis"
        #trigger sentence
        trigger_sentence = 'The science of Statistics is one thing I am amazed by.'

    elif pos == "dl":#-------------------------------------------------------------DL

        trigger_word = 'Artificial Intelligence and Supervised Learning'
        #trigger expertises
        trigger_expertise_1="optimization"
        trigger_expertise_2="forecast"
        trigger_expertise_3="analysis"
        #trigger sentence
        trigger_sentence = 'The science of Statistics is one thing I am amazed by.'

    else:#--------------------------------------------------------------------CUSTOM   
        #trigger word
        #---> {trigger_word} has been on the surface of my work as a {matched_role} more than ever before.
        #---> since then I have been involved in numerous projects involving {trigger_word}::: ''')
        trigger_word = 'Analizing Data, Buidling python programs, EDA'
        #trigger expertises
        #---> opportunities where I can use my {trigger_expertise_1}, {trigger_expertise_2}, and {trigger_expertise_3} expertises
        trigger_expertise_1="optimization"
        trigger_expertise_2="forecast"
        trigger_expertise_3="analysis"
        #trigger sentence
        #---> responsability of this job: ... I've noticed you are looking for a {position} to {trigger_sentence}::: ''')
        trigger_sentence = 'The science of Statistics is one thing I am amazed by.'


    #li_res = input('LINKED IN COVER')


    hirm_name = 'Hiring Manager'
    position_learn_from ="Linked-In"
    company_name= input("\n\n & --------------- > Company Name : ")



    #EDIT SCRIPT SAY ONE OF MY MOST similiar projects then this and its description


    #-----------------------------------------------------------------------------------------------------------------
    #-------------------------------------------------Cover-Letter----------------------------------------------------------------
    #-----------------------------------------------------------------------------------------------------------------
    #-----------------------------------------------------------------------------------------------------#### 1
    #### 1 Give a killer headline: When in doubt, use: Trgger word + Adjective(size, texture, shape , feel) +keyword +Promise
    killer_Headline = f"""I was excited to find this job opportunity is still open! I can assure you that [[ ]] 
    has been the work I have done the past years. I am very excited to continue to do [[ ]], and I am confident that the role 
    {matched_role} at {company_name} can be taken in my hands, [[ ]]."""
    #-----\
    #-----/
    #-----------------------------------------------------------------------------------------------------#### 2
    #### 2 Personalized_Greetings: Research the name of HR manager or use Dear Sir/Madam
    Personalized_Greetings= f'''Dear Sir, I am very pleased to work in my application for the {matched_role},thank you 
    for spending time reading my letter. I have lots of passion for what I do, working with [[ ]] and specifically [[]]
    have left a unique perspective that I whish nothing more than to keep building upon and to bring value with my [[]] 
    skills. I feel so grateful that this universe has been letting me keep buidling on my [[]] career, as you know 
    anything and almost everything can happen from one day to the other. I am willing to start a new position and put
    myself as a student as I believe that in {company_name} I have lots to learn. I know that my 4 years of experience working
    mostly as a [[ ]] has lots of value.

    '''

    #-----\
    #-----/
    #-----------------------------------------------------------------------------------------------------#### 3
    #### 3 Introduce Yourself:Open Strong and compliment the company
    P1_Introduce_yourself = f'''My name is Yeriko Vargas, and after graduating from my masters in applied stats I had 
    the chance to jump righ in the industry, my every first project was a tornado as I was taken a full FCA optimization
    project by meself, with a support from an advisor. coming from Statistcis to strtaight Deep LEarnin project was not easy but
    this just shows how I am able to break apart big challenges. 
    
     '''
    #-----\
    #-----/

    #-----\
    #-----/
    #-----------------------------------------------------------------------------------------------------#### 4
    #### 4 Why you are a great fit for the company:Write a short Summaryof your career and skills, and tailor it to fit to company
    P2_Why_you_great_fit= f'''The opportunity of {position} at {company_name}, would be best if taken into my hands. 
    As {matched_role} I have had the opportunity to [[ ]] .I am now working as a freelance at CAI and I recently 
    started to {cai_e1} which [[ ]] I have some experience doing more of [[ ]] Also the way I [[ ]]] 
    at [[[ ]]] started expandng in my toolbelt with [[ ]]. The [[ ]] I experienced so much at FORD, 
    where I {ford_e2} and [[ ]] {added_e3}
    '''
    #-----\
    #-----/

    #-----\
    #-----/
    #### 5 Why the COMPANY is  agreat fit for you:What excites you about working there? what do you want to learn there?
    P3_Why_company_great_fit = f'''I have experienced {trigger_word} work in detailed and I have gained so much understanding 
    around this topic, and I can assure you I have what it takes to help as {matched_role}. '''

    #### 6 Finish strong and stay in touch:Repeat why you are a great fit, explain how and when you are going to contact them
    Closing_stay_in_touch = f'''Thank you I believe that there is so much I could passionately talk 
    more in detail about my experiences. I think that the purpose of the {position} at {company_name} closely matches my 
    careergoals as well. I really would be so happy and excited to have the chance to speak with 
    someone and learn more about this position, as I want to find a deal to work here at {company_name} as for the {position}
    position'''

    #### 7 All Best...
    All_the_best_name = f'''All best,
    Yeriko Vargas'''

    killer_Headline=killer_Headline.replace('\n', '')
    Personalized_Greetings = Personalized_Greetings.replace('\n', '')
    P1_Introduce_yourself = P1_Introduce_yourself.replace('\n', '')
    P2_Why_you_great_fit = P2_Why_you_great_fit.replace('\n', '')
    P3_Why_company_great_fit =P3_Why_company_great_fit.replace('\n', '')
    Closing_stay_in_touch = Closing_stay_in_touch.replace('\n', '')

    #-----------------------------------------------------------------------------------------------------------------

    #Write txt document
    cv_text =f'''Yeriko Vargas 
    {matched_role} 
    ygvargas93@gmail.com | in/yeriko-vargas/ | twitter.com/vargasyeriko | (646)-771-6111

    {written_date}

    {company_name}
    re: {position}

    {killer_Headline}

    {Personalized_Greetings} {P1_Introduce_yourself}

    {P2_Why_you_great_fit}

    {P3_Why_company_great_fit}

    {Closing_stay_in_touch}

    {All_the_best_name}'''

    # Write file
    document.add_paragraph(f'{cv_text}')




    cv_text_chat_gpt =f'''
    HEY start here reading my CV until the end :
    
    
    Yeriko Vargas 
    {matched_role} 
    ygvargas93@gmail.com | in/yeriko-vargas/ | twitter.com/vargasyeriko | (646)-771-6111

    {written_date}

    {company_name}
    re: {position}

    {killer_Headline}

    {Personalized_Greetings} {P1_Introduce_yourself}

    {P2_Why_you_great_fit}

    {P3_Why_company_great_fit}

    {Closing_stay_in_touch}

    {All_the_best_name}
    
    
    _________________________ END OF COVER LETTER... 
    
    * NEXT STATEMENT ARE MY INSTRUCTIONS FOR YOU :::
    
    (1) : NOW you CHAT GPT change a little bit in my CV very nicely (flowy structure making sure to have by itself the CV in a 
    first markdown )matching the job description. Also mathc most important skills in experiences also found below.
    . make sure
    you use a very neutral tone. do not adjust so much the way I say things, but greatly improve my chances
    to get a match from the company I am applying to ability to match job descriptions with applicants. I do not 
    want my application to be denied so please do your best to check in with the most important words we need to 
    highlight you know. ALSO AT THE VERY END in a separate markdown provide me with those keywords you found,
    that will greatly improve the odds of me getting the position. 
    :::::
    ------- END OF INSTRUCTIONS 
    
     ::: *** NEXT STATEMENT is Structue of CV *REMEMBER I NEED A FORMAT OF A COVER LETTER with the FOLLOWING POINTS without 
     mentioning the titles, dont be akeward write as a human , a letter to a nother human so they see my potential (in a MARKDOWN):
    <<<<<<
    Yeriko Vargas 
    {matched_role} 
    ygvargas93@gmail.com | in/yeriko-vargas/ | twitter.com/vargasyeriko | (646)-771-6111

    {written_date}

    {company_name}
    re: {position}

    #### 1 Give a killer headline: When in doubt, use: Trgger word + Adjective(size, texture, shape , feel) +keyword +Promise

    #### 2 Personalized_Greetings: Research the name of HR manager or use Dear Sir/Madam

    #### 3 Introduce Yourself:Open Strong and compliment the company

    #### 4 Why you are a great fit for the company:Write a short Summaryof your career and skills, and tailor it to fit to company

    #### 5 Why the COMPANY is  agreat fit for you:What excites you about working there? what do you want to learn there?

    #### 6 Finish strong and stay in touch:Repeat why you are a great fit, explain how and when you are going to contact them
    
    #### 7 All Best...
    All best,
    Yeriko Vargas
    >>>>>>
    
    *** NEXT STATEMENT is JOB DESCRIPTION I AM APPLYING TO :

    {job_desc}

    REMEMBER ! hey do not exaggerate things, make it relatable but make sure you 
    understand I am not lying right here , just relate 
    '''


        # Save File
        #document.save(f'test.docx')
    pyperclip.copy(cv_text_chat_gpt)


    import shutil
    print("\nYOUR CV is BEING WRITTEN and NOW GO to !!!! CHAT GPT and paste !!! return to FINISH CV... \n")

# 4 --<<EXPORTS>>
    direc_exp = '/Users/yerik/_apple_source/all_data/temp/_0ds/_0_ds_Employment/'
    direc_arch_cv = '/Users/yerik/_apple_source/d_OUT/zARCHIVE/_0ds/cover_letters/'
    direc_arch_res = '/Users/yerik/_apple_source/d_OUT/zARCHIVE/_0ds/resumes/'
#'/Users/yerik/_apple_source/d_OUT/_0ds/apply_now/'

    
    document.save(f'{direc_exp}Yeriko Vargas - Cover Letter.docx')#WORD
    # OPEN CV
    subprocess.Popen(['open', f'{direc_exp}Yeriko Vargas - Cover Letter.docx'])
    # EDIT then come back and 
    input('COPIED :CV_statement now go CHAT GPT :: MODIFY CV and SAVE to DESKTOP by -> ENTER \n')
    
    # (1)2PDF
    # convert(f"{direc_exp}Yeriko Vargas - Cover Letter.docx",
    #         f"{direc_arch_cv}Yeriko Vargas-Cover_Letter - {company_name}_{position}-{date}.pdf" ) #pdf
    # input("\nENTER to PASTE CV to DESKTOP \n")
    convert(f"{direc_exp}Yeriko Vargas - Cover Letter.docx",
            f"/Users/yerik/Desktop/Yeriko Vargas - Cover Letter.pdf") 


    #####------------------------------------------ RESUME

    res_stat = f''' NOW that you know more about my background and the position, further check my analysis for 
    my relatable expereinces ::>>>: {schema} :<<<:: 
    
    *NOW write 4-5 bullet points for each experience you mentioned in the COVER LETTER. (very organized, clean and CLEAR 
    powerful statements.
    *Also write at the beggining a 3 line summary to maximize the chances I get the job. Be aware of the way 
    that companies select resumes and write those bullet point including the most important keywords from the job desc.
    * DONT WRITE DUPLICATE EXPERIENCES - check the CODE of the experience and summarize same experiences together under one job experience

    '''

    pyperclip.copy(res_stat)
    # 1 COPY MASTER RESUME TO TEMP LOCATION
    master_resume_path = '/Users/yerik/_apple_source/all_data/_0ds/0_DS_Resume_Yeriko_Vargas.docx'
    temp_res_loc_path = f'{direc_exp}/Yeriko Vargas - Resume.docx'
    shutil.copy2(master_resume_path, temp_res_loc_path)

    input('\n ... <E> ::: WAIT COPYING MASTER RESUME to TEMP 5 sec then enter to MODIFY RESUME ... \n')
    # 2 OPEN TEMP resume modify it and save it to ARCH and Desktop in a PDF
    subprocess.Popen(['open',f'{direc_exp}/Yeriko Vargas - Resume.docx'])

    input('COPIED :RESUME_statement / go CHAT GPT :: MODIFY RESUME & -> ENTER to (1) send to DESKTOP and (2) zARCHIVE  \n')
        # (1)2PDF
    convert(f'{direc_exp}Yeriko Vargas - Resume.docx',
            f"{direc_arch_res}appRES_{company_name}_{date}_RESUME_{company_name}_{position}-.pdf" ) #pdf


    convert(f'{direc_exp}Yeriko Vargas - Resume.docx',
            f"/Users/yerik/Desktop/Yeriko Vargas - Resume.pdf")
   
    #####------------------------------------------ 
    
    # (2)zARCHIVE
    # direc_arch = '/Users/yerik/_apple_source/d_OUT/zARCHIVE/_0ds/apply_now/cvs/'
    # my_cv = docx.Document(f'{direc_exp}Yeriko Vargas-Cover_Letter - {company_name}_{position}-{date}.docx') # read modified
    # my_cv.save(f'{direc_arch}Yeriko Vargas-Cover_Letter - {company_name}_{position}-{date}.docx')#WORD
    #print("SAVED TO ARCHIVE")

    # ARCHIVE
    # OPEN CV again 
    add_to_arch = f'SCHEMA :: \n\n {schema} \n\n {job_desc}'
    pyperclip.copy(add_to_arch)
    print('>>>>>>>>> after archived process is complete ::: ONE LAST ENTER TO BE DONE >>>>>>>>>>> ')
    
    subprocess.Popen(['open', f'{direc_exp}Yeriko Vargas - Cover Letter.docx'])
    input("::: ATTN -> ///  JD & SCHEMA /// COPIED ::: ENTER > TO OPEN DOC , paste to ARCHIVE and ENTER ....")
    convert(f"{direc_exp}Yeriko Vargas - Cover Letter.docx",
            f"{direc_arch_cv}appCV_{company_name}_{date}_Yeriko Vargas-Cover_Letter - {company_name}_{position}-.pdf" )
    ##############################################################


    ############################################################## 2 TOGETHER
    pdf_path1 = f'/Users/yerik/Desktop/Yeriko Vargas - Cover Letter.pdf'
    pdf_path2 = f'/Users/yerik/Desktop/Yeriko Vargas - Resume.pdf'

    from PyPDF2 import PdfReader, PdfWriter
    
    def combine_pdfs(pdf_path1, pdf_path2, output_pdf_path):
        # Create a PDF writer object
        writer = PdfWriter()
    
        # Open and read the first PDF file
        reader1 = PdfReader(pdf_path1)
        for page in reader1.pages:
            writer.add_page(page)
    
        # Open and read the second PDF file
        reader2 = PdfReader(pdf_path2)
        for page in reader2.pages:
            writer.add_page(page)
    
        # Write out the combined PDF to a new file
        with open(output_pdf_path, 'wb') as output_pdf:
            writer.write(output_pdf)
    
    # Example usage
    
    output_pdf_path =  f'/Users/yerik/Desktop/Resume_CV - position - Yeriko Vargas.pdf'
    
    combine_pdfs(pdf_path1, pdf_path2, output_pdf_path)


    ##############################################################

    
    path_tmep_j_desc = '/Users/yerik/_apple_source/all_data/temp/_0ds/_0_ds_Employment'
    # List of text variables and their respective content for processing
    text_variables = [cai_e1 ,ford_e2,added_e3]
    keywords = ['CAI', 'FORD', 'SA', 'FCA', 'ODS']

    # Function to replace the found code in the text with ' ' (a space) before saving
    def replace_code_and_save(texts, keywords):
        for text in texts:
            # Find the first keyword that appears in the text
            found_code = next((keyword for keyword in keywords if keyword in text), None)
            if found_code:
                # Replace the found code in the text with ' '
                modified_text = text.replace(found_code, ' ')
                modified_text = modified_text.replace('job description matching jd', ' ')
                modified_text = modified_text.replace('* [', ' ')
                modified_text = modified_text.replace('] ->', ' ')




                
                # Define the filename based on the found code
                filename = f"{path_tmep_j_desc}/{found_code}_custom_task_desk.txt"
                # Save the modified text to a file with the filename based on the found code
                with open(filename, 'w') as file:
                    file.write(modified_text)
                print(f"File saved with code replaced: {filename}")
    
    # Run the function to replace the codes in each text and save the files again
    replace_code_and_save(text_variables, keywords)





    
    return document

# What happens : CV is created in temp, also jd.text <_0_ds_Employment>
# then after modified it goes to zARCHIVE and a copy PDF in desktop to submit (missing one in d_OUT with other docsif needed)
# then Archived CV opens to save a further jd with it and it saves back to zARCHIVED (2 CV saved, one alone and second with info)